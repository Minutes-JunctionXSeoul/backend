import docx
import time
import numpy
import os
from datetime import datetime

from .log_summary import extract_keywords
from .calendar_utils import df_to_ics
from .entity_extractor import extract_entities

def make_docx(client,documents):
    
    keyword_sentence_dict = extract_keywords(client,documents)
    df = extract_entities(client, documents)

    dt = datetime.now()

    year = dt.strftime("%y")
    month = dt.strftime("%m")
    day = dt.strftime("%d")
    
    company_name = "X company"
    location = "conference room"
    
    date = year + '.' + month + '.' +  day + '.'
    
    hour = dt.strftime("%H")
    mins = dt.strftime("%M")
    
    time = hour + ":" + mins
    
    
    
    doc = docx.Document(os.path.abspath("api/forms/form1.docx"))

    para = doc.paragraphs
    table = doc.tables
        
    #개요
    table[0].cell(0,1).text = location
    table[0].cell(1,1).text = date
    table[0].cell(2,1).text = time
    
    #keyword
    for key in keyword_sentence_dict.keys():
        keyword = para[2].add_run(key + ' / ')
        keyword.bold = True
        keyword.font.size = docx.shared.Pt(12)
    
    #summary
    keyword_in_sentence = list(keyword_sentence_dict.values())
    keywords = list(keyword_sentence_dict.keys())
    for i in range(len(documents)):
        if documents[i] in keyword_in_sentence:
            keyword_in_sentence = numpy.array(keyword_in_sentence)
            idx = numpy.where(keyword_in_sentence == documents[i])[0]
            write_kw = ""
            for k in idx:
                write_kw += keywords[k]+" / "
            keyword = para[7].add_run("\n" + write_kw[:-2] + "\n")
            keyword.bold = True
            keyword.font.size = docx.shared.Pt(13)
        sentence = para[7].add_run(documents[i]+"\n")

    #todo-list
    todo_df = df_to_ics(df)  
    if not todo_df.empty:
        for i in range(len(todo_df)):
            current = todo_df.iloc[i]
            start_date = current['Start Date']
            start_time = current['Start Time']
            end_date = current['End Date']
            end_time = current['End Time']
            subject = current['Subject']
            description = current['Description']
            table[1].add_row()
            table[1].cell(i+1,0).text = subject
            table[1].cell(i+1,1).text = start_date + " " + start_time
            table[1].cell(i+1,2).text = end_date + end_time
            table[1].cell(i+1,3).text = description
    else:
        for i in range(4):
            table[1].cell(1,i).text = "Empty"
        
    
    file_name = month + day + "_" + hour + mins + '_meeting minutes.docx'
    doc.save(file_name)
    return file_name
    
